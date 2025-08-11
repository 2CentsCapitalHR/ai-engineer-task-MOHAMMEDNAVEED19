from docx import Document


def annotate_docx_with_issues(input_path, issues, output_path) -> bool:
    try:
        doc = Document(input_path)
        paras = [p for p in doc.paragraphs]
        new_doc = Document()
        for i, p in enumerate(paras):
            new_doc.add_paragraph(p.text)
            for issue in [it for it in issues if it.get("para_index") == i]:
                note_text = f"<< REVIEW [{issue.get('severity','Medium')}]: {issue.get('issue')} >>"
                if issue.get('suggestion'):
                    note_text += f"\nSuggested rewrite: {issue.get('suggestion')}"
                if issue.get('citations'):
                    note_text += f"\nReferences: {', '.join(issue.get('citations', []))}"
                note = new_doc.add_paragraph(note_text)
                # keep simple formatting to avoid complex docx color handling
                note.runs[0].italic = True
        new_doc.add_page_break()
        new_doc.add_heading("Review Summary", level=1)
        for idx, issue in enumerate(issues, start=1):
            new_doc.add_paragraph(f"{idx}. Paragraph {issue.get('para_index')}: {issue.get('issue')}", style='List Bullet')
            new_doc.add_paragraph(f"   Severity: {issue.get('severity')}")
            if issue.get('suggestion'):
                new_doc.add_paragraph(f"   Suggestion: {issue.get('suggestion')}")
            if issue.get('citations'):
                new_doc.add_paragraph(f"   References: {', '.join(issue.get('citations', []))}")
            new_doc.add_paragraph("")
        new_doc.save(output_path)
        return True
    except Exception as e:
        print(f"[ERROR] Failed to annotate document {input_path.name}: {e}")
        return False
